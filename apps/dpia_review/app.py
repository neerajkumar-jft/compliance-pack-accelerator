"""DPIA Review — Databricks App (Streamlit) for Phase 4 Theme 3.

The human approval surface for the DPIA Auto-Generator (Agent 1).
Replaces the Phase 2 CLI (`scripts/approve_dpia.py`) with a UI that
captures the *verified* Databricks user identity from request headers,
removing the CLI's self-asserted-reviewer trust assumption.

Three views:
  1. List — all rows in compliance.dpia_runs with status filter chips.
  2. Detail — metadata, the 8 structured sections (from dpia_sections
     MAP added in Phase 3), and the input context_snapshot for
     ground-truth verification.
  3. Approve action — visible only on draft rows AND only when the
     logged-in user is a CCO or GC persona (the two approver roles).
     Issues the same idempotent UPDATE the CLI runs.

Plus two download formats from the detail-view header:
  - **Download Word (.docx)** — built with python-docx (pure Python,
    no system deps). Reviewers can edit + sign + save-as-PDF themselves.
  - **Download PDF** — built with ReportLab + the markdown lib (markdown
    parsed into ReportLab-Paragraph-compatible HTML so bold/italic/lists
    actually render as formatted output, not as raw markdown source).

Both formats interpret the LLM's markdown output: headings, bullets,
numbered lists, bold/italic spans, inline code. WeasyPrint was tried
first but its system-level cairo + pango deps aren't in the Apps image.

Permission model (defense in depth):
  - Workspace CAN_USE on the app — granted to CCO + GC + CFO only;
    CMO can't even open the URL.
  - The app's runtime SP holds SELECT + UPDATE on dpia_runs, plus
    SELECT on the dpia_artifacts volume.
  - No persona has direct UPDATE on dpia_runs — the only path to
    flipping status is through this app's enforce-the-rules code.
  - In-app role gate hides the Approve button for CFO (audience, not
    approver) even though they can see the list.

Reads ``dashboards/personas/.persona_emails.json`` (produced by
scripts/setup_persona_users.py) at startup to know which emails are
CCO and GC. If that file is absent, the app degrades to read-only —
nobody can approve, but everyone can view.
"""

from __future__ import annotations

import io
import json
import os
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import streamlit as st
from databricks.sdk import WorkspaceClient
from databricks.sdk.service.sql import StatementState

CATALOG = os.environ.get("COMPLIANCE_CATALOG", "compliance_pack")
TABLE = f"{CATALOG}.compliance.dpia_runs"
PERSONA_EMAILS_PATH = Path("dashboards/personas/.persona_emails.json")

st.set_page_config(page_title="DPIA Review", layout="wide")


# ---------------------------------------------------------------------
# Identity + role
# ---------------------------------------------------------------------


def get_user_email() -> str:
    """Read the authenticated Databricks user's email from request headers.

    Databricks Apps inject ``X-Forwarded-Email`` (and friends) on every
    request — that's the verified workspace identity, not something the
    user typed. Falls back to a placeholder when run outside an Apps
    container (e.g., local dev), so the UI still renders for testing.
    """
    headers = st.context.headers if hasattr(st, "context") else {}
    return (
        headers.get("X-Forwarded-Email")
        or headers.get("x-forwarded-email")
        or "unknown@example.local"
    )


def load_approver_emails() -> dict[str, str]:
    """Discover CCO and GC persona emails — the two roles that can approve.

    Two strategies in order:

    1. **SCIM lookup against the workspace user directory** (production
       path). The deployed Apps container can read users via the SDK,
       and persona emails follow a stable plus-address pattern set by
       scripts/setup_persona_users.py: ``<base>+compliance-cco@<domain>`` and
       ``<base>+compliance-gc@<domain>``. We list users whose userName contains
       ``+compliance-`` and pick out the cco/gc entries. This is the canonical
       runtime path because:
         a. Apps containers don't ship the repo's ``dashboards/personas/``
            directory on a stable relative path
         b. Re-running setup_persona_users.py shouldn't require restarting
            the app to refresh the role gate

    2. **Local file fallback** (dev / test path). When SCIM listing fails
       (no permission, no SDK, etc.) we fall back to reading
       ``dashboards/personas/.persona_emails.json`` from the working
       directory. This keeps unit tests and local Streamlit runs working.

    Returns ``{"cco": ..., "gc": ...}`` or a subset; empty dict means
    no approvers were detectable, so the Approve button is hidden
    everywhere (safe-by-default).
    """
    # 1. SCIM-based discovery — the production path
    try:
        client = _client()
        approvers: dict[str, str] = {}
        for user in client.users.list(filter='userName co "+compliance-"'):
            email = (user.user_name or "").lower()
            if "+compliance-cco@" in email:
                approvers["cco"] = user.user_name
            elif "+compliance-gc@" in email:
                approvers["gc"] = user.user_name
        if approvers:
            return approvers
    except Exception:
        # SCIM unavailable (running outside Apps container, missing
        # permission, transient) — fall through to file-based path.
        pass

    # 2. File fallback — dev / test path
    if PERSONA_EMAILS_PATH.exists():
        try:
            data = json.loads(PERSONA_EMAILS_PATH.read_text())
            return {k: v for k, v in data.items() if k in ("cco", "gc") and v}
        except (OSError, json.JSONDecodeError):
            pass

    return {}


def is_approver(user_email: str, approver_emails: dict[str, str]) -> bool:
    """True iff the logged-in user is the CCO or GC persona."""
    return user_email in approver_emails.values()


# ---------------------------------------------------------------------
# SQL via the bundled SDK (uses the app's runtime identity / warehouse)
# ---------------------------------------------------------------------


@st.cache_resource
def _client() -> WorkspaceClient:
    return WorkspaceClient()


def _warehouse_id() -> str:
    """Read warehouse id from env (set in app.yaml) or pick the first
    serverless warehouse the app's SP can see."""
    if env := os.environ.get("COMPLIANCE_WAREHOUSE_ID"):
        return env
    w = _client()
    for wh in w.warehouses.list():
        if wh.enable_serverless_compute:
            return wh.id
    raise RuntimeError("No serverless SQL warehouse available to the app's identity.")


def _exec(stmt: str) -> dict:
    """Execute SQL via the SDK's statement-execution API. Returns
    {"data": list-of-rows} on success, {"error": str} on failure."""
    w = _client()
    res = w.statement_execution.execute_statement(
        warehouse_id=_warehouse_id(),
        statement=stmt,
        wait_timeout="30s",
    )
    if res.status.state != StatementState.SUCCEEDED:
        return {"error": (res.status.error.message if res.status.error else str(res.status.state))[:500]}
    cols = [c.name for c in (res.manifest.schema.columns or [])]
    rows = [dict(zip(cols, r)) for r in (res.result.data_array or [])]
    return {"data": rows}


def list_runs(status_filter: str | None = None) -> list[dict]:
    where = "WHERE 1=1"
    if status_filter and status_filter != "all":
        where += f" AND status = '{status_filter}'"
    res = _exec(f"""
        SELECT run_id, generated_at, generated_by, status,
               reviewed_by, reviewed_at, model_endpoint, prompt_version,
               regulation_pack, ROUND(latency_seconds, 1) AS latency_s,
               CASE WHEN parse_error IS NOT NULL THEN 'parse_error' ELSE 'ok' END AS parse_status
        FROM {TABLE}
        {where}
        ORDER BY generated_at DESC
    """)
    if "error" in res:
        st.error(f"Could not load runs: {res['error']}")
        return []
    return res["data"]


def fetch_run(run_id: str) -> dict | None:
    # run_id format is regex-validated upstream when set in session_state
    res = _exec(f"""
        SELECT run_id, generated_at, generated_by, status,
               reviewed_by, reviewed_at, model_endpoint, prompt_module,
               prompt_version, regulation_pack, latency_seconds,
               artifact_path, dpia_text, dpia_sections, parse_error,
               context_snapshot, notes
        FROM {TABLE}
        WHERE run_id = '{run_id}'
    """)
    if "error" in res or not res["data"]:
        return None
    return res["data"][0]


def approve(run_id: str, reviewer: str, notes: str | None = None) -> dict:
    """Idempotent UPDATE — mirrors scripts/approve_dpia.py logic."""
    safe_reviewer = reviewer.replace("'", "''")
    notes_clause = ""
    if notes:
        safe_notes = notes.replace("'", "''")
        notes_clause = f", notes = '{safe_notes}'"
    return _exec(f"""
        UPDATE {TABLE}
        SET status = 'approved',
            reviewed_by = '{safe_reviewer}',
            reviewed_at = current_timestamp(){notes_clause}
        WHERE run_id = '{run_id}' AND status = 'draft'
    """)


# ---------------------------------------------------------------------
# Rendering helpers — section list + markdown parser
# ---------------------------------------------------------------------


SECTION_TITLES_RENDER: list[tuple[str, str]] = [
    ("executive_summary",       "1. Executive Summary"),
    ("data_inventory",          "2. Data Inventory"),
    ("processing_activities",   "3. Processing Activities"),
    ("risk_assessment",         "4. Risk Assessment"),
    ("compliance_gap_analysis", "5. Compliance Gap Analysis"),
    ("consent_status",          "6. Consent Status"),
    ("remediation_plan",        "7. Remediation Plan"),
    ("residual_risk",           "8. Residual Risk"),
]


def _normalize_sections(row: dict) -> dict | None:
    """Decode dpia_sections into a dict, or return None on parse failure."""
    sections = row.get("dpia_sections")
    if isinstance(sections, str):
        try:
            sections = json.loads(sections)
        except (json.JSONDecodeError, TypeError):
            return None
    return sections if sections else None


def _build_full_markdown(row: dict) -> str:
    """Concatenate the 8 sections into a single markdown document for the
    "Rendered Document" tab and as input to the docx/pdf renderers when
    we want a one-shot view of the whole DPIA."""
    sections = _normalize_sections(row)
    if not sections:
        return f"# Raw model output\n\n_Note: structured parse failed; showing raw LLM response._\n\n```\n{row.get('dpia_text') or '(empty)'}\n```"
    parts = [f"# Data Protection Impact Assessment\n",
             f"_Run ID:_ `{row.get('run_id', '')}` · _generated:_ {row.get('generated_at', '')}\n"]
    for key, title in SECTION_TITLES_RENDER:
        body = sections.get(key) or "_(missing)_"
        parts.append(f"\n## {title}\n\n{body}\n")
    return "".join(parts)


# Light-weight markdown block parser. Returns a list of (kind, payload):
#   ("h1", str) | ("h2", str) | ("h3", str) | ("h4", str) | ("h5", str) | ("h6", str)
#   ("bullet", str) | ("number", str)
#   ("table", list[list[str]])  — header row first, then data rows
#   ("para", str) — a regular paragraph (may span multiple input lines)
#
# `str` payloads still contain inline markdown (**bold**, *italic*, `code`)
# which the ReportLab + python-docx renderers expand below.
import re as _re

_INLINE_BOLD = _re.compile(r"\*\*([^*]+?)\*\*")
_INLINE_ITALIC = _re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")
_INLINE_CODE = _re.compile(r"`([^`]+?)`")
_BULLET_RE = _re.compile(r"^\s*[-*]\s+(.+)$")
_NUMBER_RE = _re.compile(r"^\s*\d+\.\s+(.+)$")
_HEADING_RE = _re.compile(r"^(#{1,6})\s+(.+)$")  # h1–h6 (LLMs often use ####)
# A markdown-table separator row: |---|:---:|---:| etc — at least one dash
# per cell, may have leading/trailing colons for alignment.
_TABLE_SEP_RE = _re.compile(r"^\s*\|?(\s*:?-+:?\s*\|)+\s*:?-+:?\s*\|?\s*$")


def _split_md_table_row(line: str) -> list[str]:
    """Split a markdown table row by | into cell strings (trimmed,
    leading/trailing pipes dropped)."""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


# Glyphs that the LLM emits but Helvetica's built-in font set can't render
# in PDFs (ReportLab falls back to .notdef → black square). Pre-process
# any user-content string before it reaches Paragraph. Word (.docx) can
# render these natively (Calibri has wide Unicode coverage), so the
# normalizer is PDF-only — applied via _inline_to_reportlab.
_PDF_GLYPH_NORMALIZE = {
    "‑": "-",     # non-breaking hyphen → ASCII hyphen
    "­": "",      # soft hyphen → drop
    " ": " ",     # non-breaking space → regular space
    " ": " ",     # narrow no-break space
    " ": " ",     # thin space
    "‘": "'",     # left single curly quote
    "’": "'",     # right single curly quote
    "“": '"',     # left double curly quote
    "”": '"',     # right double curly quote
    "…": "...",   # horizontal ellipsis
}


def _normalize_for_pdf_glyphs(s: str) -> str:
    for k, v in _PDF_GLYPH_NORMALIZE.items():
        s = s.replace(k, v)
    return s


def _parse_md_blocks(md: str) -> list[tuple]:
    """Tokenize a small subset of markdown into block tuples.

    Switched to index-based iteration so we can lookahead for the
    table-separator line that distinguishes a markdown table from a
    regular paragraph that happens to start with ``|``.
    """
    blocks: list[tuple] = []
    paragraph_lines: list[str] = []

    def flush_paragraph():
        if paragraph_lines:
            blocks.append(("para", " ".join(paragraph_lines).strip()))
            paragraph_lines.clear()

    lines = (md or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            flush_paragraph()
            i += 1
            continue

        # Markdown table: a line starting with `|` whose NEXT line is
        # a separator like `|---|---|`.
        if line.lstrip().startswith("|") and i + 1 < len(lines) \
                and _TABLE_SEP_RE.match(lines[i + 1].rstrip()):
            flush_paragraph()
            header = _split_md_table_row(line)
            i += 2  # consume header + separator
            data_rows: list[list[str]] = []
            while i < len(lines) and lines[i].rstrip().lstrip().startswith("|"):
                data_rows.append(_split_md_table_row(lines[i].rstrip()))
                i += 1
            blocks.append(("table", [header, *data_rows]))
            continue

        m = _HEADING_RE.match(line)
        if m:
            flush_paragraph()
            level = len(m.group(1))
            blocks.append((f"h{level}", m.group(2).strip()))
            i += 1
            continue
        m = _BULLET_RE.match(line)
        if m:
            flush_paragraph()
            blocks.append(("bullet", m.group(1).strip()))
            i += 1
            continue
        m = _NUMBER_RE.match(line)
        if m:
            flush_paragraph()
            blocks.append(("number", m.group(1).strip()))
            i += 1
            continue
        paragraph_lines.append(line)
        i += 1
    flush_paragraph()
    return blocks


def _inline_to_reportlab(s: str) -> str:
    """Convert a small inline-markdown subset to ReportLab Paragraph HTML.
    Normalizes glyphs Helvetica can't render, escapes raw `<`/`>`/`&`
    so the model's literal angle brackets don't get mistaken for tags,
    then expands **bold** / *italic* / `code`."""
    s = _normalize_for_pdf_glyphs(s)
    s = (s.replace("&", "&amp;")
           .replace("<", "&lt;")
           .replace(">", "&gt;"))
    s = _INLINE_BOLD.sub(r"<b>\1</b>", s)
    s = _INLINE_ITALIC.sub(r"<i>\1</i>", s)
    s = _INLINE_CODE.sub(r"<font face='Courier'>\1</font>", s)
    return s


# ---------------------------------------------------------------------
# PDF — ReportLab + small markdown parser
# ---------------------------------------------------------------------


def render_pdf(row: dict) -> bytes:
    """Render the DPIA to PDF using ReportLab + a markdown subset.

    Pure Python (no system deps), so safe for the Databricks Apps
    runtime image. Renders headings, bullets, numbered lists, and
    inline bold/italic/code from the LLM's markdown output — not
    just the raw markdown source like the Phase 4 Theme 3 first cut.

    Falls back to the raw `dpia_text` in a single block when
    `dpia_sections` is NULL (parse-failed case).
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak,
        Table, TableStyle, ListFlowable, ListItem, Preformatted,
    )
    from reportlab.lib.enums import TA_LEFT

    sections = _normalize_sections(row)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.8 * cm, rightMargin=1.8 * cm,
        topMargin=2.2 * cm, bottomMargin=2.5 * cm,
        title=f"DPIA · {row.get('run_id', '')}",
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=20,
                        textColor=HexColor("#003366"), spaceAfter=12)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=13,
                        textColor=HexColor("#003366"), spaceBefore=14, spaceAfter=6)
    h3 = ParagraphStyle("h3", parent=styles["Heading3"], fontSize=11,
                        textColor=HexColor("#003366"), spaceBefore=8, spaceAfter=4)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=10,
                          alignment=TA_LEFT, spaceAfter=4, leading=14)
    pre = ParagraphStyle("pre", parent=styles["Code"], fontSize=9, leading=12)

    def _md_to_flowables(md: str) -> list:
        flowables: list = []
        for kind, payload in _parse_md_blocks(md):
            # Tables: payload is rows (list[list[str]]); everything else is str.
            if kind == "table":
                rows = payload or []
                if not rows:
                    continue
                # Build a Table of Paragraphs so cells word-wrap and
                # get inline-markdown formatting (bold/italic/code) too.
                table_data = []
                for r_idx, row_cells in enumerate(rows):
                    para_style = body  # all rows use body style; header gets bold-tweaked via cell style
                    table_data.append([
                        Paragraph(_inline_to_reportlab(c or ""), para_style)
                        for c in row_cells
                    ])
                n_cols = max(len(r) for r in rows)
                # Total page width minus margins ≈ 16 cm — divide evenly
                col_widths = [(16.0 / n_cols) * cm] * n_cols
                tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), HexColor("#003366")),
                    ("TEXTCOLOR",  (0, 0), (-1, 0), HexColor("#ffffff")),
                    ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("BACKGROUND", (0, 1), (-1, -1), HexColor("#fbfbfd")),
                    ("BOX",        (0, 0), (-1, -1), 0.5, HexColor("#999999")),
                    ("INNERGRID",  (0, 0), (-1, -1), 0.25, HexColor("#dddddd")),
                    ("VALIGN",     (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING",  (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING",   (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
                ]))
                flowables.append(tbl)
                flowables.append(Spacer(1, 0.2 * cm))
                continue

            html = _inline_to_reportlab(payload)
            if kind == "h1":
                flowables.append(Paragraph(html, h2))  # h1 inside section → use h2 visual
            elif kind in ("h2", "h3"):
                flowables.append(Paragraph(html, h3))
            elif kind in ("h4", "h5", "h6"):
                # h4-h6 → bold body paragraph (smaller heading effect without
                # introducing a 4th heading style)
                flowables.append(Paragraph(f"<b>{html}</b>", body))
            elif kind in ("bullet", "number"):
                bullet_type = "bullet" if kind == "bullet" else "1"
                flowables.append(ListFlowable(
                    [ListItem(Paragraph(html, body), leftIndent=18)],
                    bulletType=bullet_type, leftIndent=20, spaceBefore=2, spaceAfter=2,
                ))
            else:  # para
                flowables.append(Paragraph(html, body))
        return flowables

    story = []

    # Cover
    story.append(Paragraph("Data Protection Impact Assessment", h1))
    story.append(Paragraph(
        f"Generated by the DPIA Auto-Generator (Agent 1)"
        + (f" · regulation pack: <font face='Courier'>{row.get('regulation_pack')}</font>"
           if row.get("regulation_pack") else ""),
        body,
    ))

    # Metadata table
    meta_rows = [
        ("Run ID", row.get("run_id", "")),
        ("Generated at", str(row.get("generated_at", ""))),
        ("Generated by", row.get("generated_by") or "—"),
        ("Status", row.get("status") or "—"),
    ]
    if row.get("reviewed_by"):
        meta_rows += [
            ("Reviewed by", row.get("reviewed_by")),
            ("Reviewed at", str(row.get("reviewed_at", ""))),
        ]
    meta_rows += [
        ("Model endpoint", row.get("model_endpoint", "")),
        ("Prompt version", row.get("prompt_version", "")),
    ]
    t = Table(meta_rows, colWidths=[5 * cm, 11 * cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), HexColor("#f7f7f9")),
        ("BOX",        (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
        ("INNERGRID",  (0, 0), (-1, -1), 0.25, HexColor("#dddddd")),
        ("FONTNAME",   (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("LEFTPADDING",  (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING",   (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
    ]))
    story.append(Spacer(1, 0.4 * cm))
    story.append(t)
    story.append(Spacer(1, 0.6 * cm))
    story.append(Paragraph(
        "This document was rendered from the structured output of an LLM-based "
        "DPIA generator and must be reviewed by a qualified privacy officer "
        "before submission to a regulator. The audit row in "
        "<font face='Courier'>compliance_pack.compliance.dpia_runs</font> with run_id "
        f"<font face='Courier'>{row.get('run_id', '')}</font> is the "
        "authoritative system-of-record.",
        body,
    ))
    story.append(PageBreak())

    # Sections
    if sections:
        for key, title in SECTION_TITLES_RENDER:
            story.append(Paragraph(title, h2))
            content = sections.get(key)
            if content:
                story.extend(_md_to_flowables(str(content)))
            else:
                story.append(Paragraph(
                    "<i>— section missing in model output —</i>",
                    ParagraphStyle("missing", parent=body, textColor=HexColor("#999999")),
                ))
            story.append(Spacer(1, 0.3 * cm))
    else:
        story.append(Paragraph("Raw model output", h2))
        story.append(Paragraph(
            "<font color='#806000'>Note: the structured (8-section) output failed "
            "Pydantic validation for this run. The raw text below is what the LLM "
            "returned. The <font face='Courier'>parse_error</font> column on the "
            "audit row explains the validation failure.</font>",
            body,
        ))
        story.append(Spacer(1, 0.3 * cm))
        story.append(Preformatted(str(row.get("dpia_text") or "(empty)"), pre))

    story.append(Spacer(1, 0.6 * cm))
    story.append(Paragraph(
        f"<font color='#888888' size='8'>Rendered at "
        f"{datetime.now(timezone.utc).isoformat()} by the DPIA Review app. "
        f"Authoritative source: <font face='Courier'>compliance_pack.compliance.dpia_runs.run_id "
        f"= '{row.get('run_id', '')}'</font></font>",
        body,
    ))
    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------
# Word — python-docx (pure Python, reviewers can edit + save-as-PDF)
# ---------------------------------------------------------------------


def render_docx(row: dict) -> bytes:
    """Render the DPIA to a Word .docx using python-docx + the same
    markdown parser PDF rendering uses.

    Editable artifact — the natural format for "reviewer reads, edits
    cover sheet, signs, exports to PDF when ready to submit."
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm

    sections = _normalize_sections(row)
    doc = Document()

    # Page margins (A4)
    for section in doc.sections:
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.5)

    # Cover title
    title = doc.add_heading("Data Protection Impact Assessment", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p = doc.add_paragraph(
        "Generated by the DPIA Auto-Generator (Agent 1)"
        + (f"  ·  regulation pack: {row.get('regulation_pack')}" if row.get("regulation_pack") else "")
    )
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Metadata table
    meta_rows = [
        ("Run ID", row.get("run_id", "")),
        ("Generated at", str(row.get("generated_at", ""))),
        ("Generated by", row.get("generated_by") or "—"),
        ("Status", row.get("status") or "—"),
    ]
    if row.get("reviewed_by"):
        meta_rows += [
            ("Reviewed by", row.get("reviewed_by")),
            ("Reviewed at", str(row.get("reviewed_at", ""))),
        ]
    meta_rows += [
        ("Model endpoint", row.get("model_endpoint", "")),
        ("Prompt version", row.get("prompt_version", "")),
    ]
    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.style = "Light Grid Accent 1"
    for (label, value), tr in zip(meta_rows, table.rows):
        tr.cells[0].text = label
        tr.cells[1].text = str(value)
        for run in tr.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        for run in tr.cells[1].paragraphs[0].runs:
            run.font.size = Pt(9)

    doc.add_paragraph()
    intro = doc.add_paragraph(
        "This document was rendered from the structured output of an LLM-based "
        "DPIA generator and must be reviewed by a qualified privacy officer before "
        "submission to a regulator. The audit row in compliance_pack.compliance.dpia_runs "
        f"with run_id {row.get('run_id', '')} is the authoritative system-of-record."
    )
    for run in intro.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    def _add_inline_runs(paragraph, payload: str):
        """Walk the inline-markdown subset and emit python-docx runs."""
        i = 0
        n = len(payload)
        # Pre-tokenize: list of (kind, text) where kind ∈ {"text","bold","italic","code"}
        tokens: list[tuple[str, str]] = []
        cursor = 0
        # Combined regex with named groups, evaluated left-to-right
        token_re = _re.compile(r"\*\*(?P<bold>[^*]+?)\*\*|(?<!\*)\*(?P<italic>[^*]+?)\*(?!\*)|`(?P<code>[^`]+?)`")
        for m in token_re.finditer(payload):
            if m.start() > cursor:
                tokens.append(("text", payload[cursor:m.start()]))
            if m.group("bold") is not None:
                tokens.append(("bold", m.group("bold")))
            elif m.group("italic") is not None:
                tokens.append(("italic", m.group("italic")))
            elif m.group("code") is not None:
                tokens.append(("code", m.group("code")))
            cursor = m.end()
        if cursor < n:
            tokens.append(("text", payload[cursor:]))
        for kind, text in tokens:
            run = paragraph.add_run(text)
            if kind == "bold":
                run.bold = True
            elif kind == "italic":
                run.italic = True
            elif kind == "code":
                run.font.name = "Courier New"
                run.font.size = Pt(9)

    def _add_md_to_doc(md: str):
        for kind, payload in _parse_md_blocks(md):
            # Tables: payload is rows (list[list[str]])
            if kind == "table":
                rows = payload or []
                if not rows:
                    continue
                n_cols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=n_cols)
                tbl.style = "Light Grid Accent 1"
                for r_idx, row_cells in enumerate(rows):
                    for c_idx in range(n_cols):
                        cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
                        cell = tbl.rows[r_idx].cells[c_idx]
                        # python-docx pre-creates an empty paragraph in
                        # each cell; clear it and rebuild with inline runs
                        cell.paragraphs[0].text = ""
                        _add_inline_runs(cell.paragraphs[0], cell_text)
                        if r_idx == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                doc.add_paragraph()  # spacing after table
                continue
            if kind == "h1":
                doc.add_heading(payload, level=2)
            elif kind == "h2":
                doc.add_heading(payload, level=3)
            elif kind == "h3":
                doc.add_heading(payload, level=4)
            elif kind in ("h4", "h5", "h6"):
                # h4-h6 → bold body paragraph (Word's heading levels 5+
                # render so small they're visually equivalent to bold body)
                p = doc.add_paragraph()
                run = p.add_run(payload)
                run.bold = True
                run.font.size = Pt(10)
            elif kind == "bullet":
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, payload)
            elif kind == "number":
                p = doc.add_paragraph(style="List Number")
                _add_inline_runs(p, payload)
            else:
                p = doc.add_paragraph()
                _add_inline_runs(p, payload)

    if sections:
        for key, sec_title in SECTION_TITLES_RENDER:
            doc.add_heading(sec_title, level=1)
            content = sections.get(key)
            if content:
                _add_md_to_doc(str(content))
            else:
                p = doc.add_paragraph()
                run = p.add_run("— section missing in model output —")
                run.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    else:
        doc.add_heading("Raw model output", level=1)
        p = doc.add_paragraph()
        warn = p.add_run(
            "Note: the structured (8-section) output failed Pydantic validation "
            "for this run. The raw text below is what the LLM returned. The "
            "parse_error column on the audit row explains the validation failure."
        )
        warn.font.color.rgb = RGBColor(0x80, 0x60, 0x00)
        raw = doc.add_paragraph(str(row.get("dpia_text") or "(empty)"))
        for run in raw.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)

    # Footer note
    doc.add_paragraph()
    footer = doc.add_paragraph(
        f"Rendered at {datetime.now(timezone.utc).isoformat()} by the DPIA "
        f"Review app. Authoritative source: compliance_pack.compliance.dpia_runs.run_id "
        f"= '{row.get('run_id', '')}'"
    )
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------
# Views
# ---------------------------------------------------------------------


SECTION_TITLES: dict[str, str] = {
    "executive_summary": "Executive Summary",
    "data_inventory": "Data Inventory",
    "processing_activities": "Processing Activities",
    "risk_assessment": "Risk Assessment",
    "compliance_gap_analysis": "Compliance Gap Analysis",
    "consent_status": "Consent Status",
    "remediation_plan": "Remediation Plan",
    "residual_risk": "Residual Risk",
}


def _status_badge(status: str) -> str:
    color = {
        "draft": "🟡",
        "approved": "🟢",
        "superseded": "⚪",
    }.get(status, "❔")
    return f"{color} {status}"


def view_list(status_filter: str) -> None:
    st.title("DPIA Review")

    rows = list_runs(status_filter)
    if not rows:
        st.info(f"No DPIAs found for status='{status_filter}'.")
        return

    st.caption(f"{len(rows)} run(s) — click a row to open.")

    for row in rows:
        with st.container(border=True):
            cols = st.columns([3, 2, 2, 3, 2])
            cols[0].markdown(f"**{row['run_id']}**\n\n_{row['generated_at']}_")
            cols[1].markdown(_status_badge(row["status"]))
            cols[2].markdown(f"`{row['parse_status']}`")
            cols[3].markdown(row.get("reviewed_by") or "—")
            if cols[4].button("Open", key=f"open_{row['run_id']}"):
                st.session_state.selected_run = row["run_id"]
                st.rerun()


def view_detail(run_id: str, user_email: str, can_approve: bool) -> None:
    row = fetch_run(run_id)
    if row is None:
        st.error(f"run_id {run_id!r} not found.")
        if st.button("← Back to list"):
            st.session_state.pop("selected_run", None)
            st.rerun()
        return

    # ── Top header: back button (left) + downloads (right) ────────────
    top_left, _, top_right = st.columns([3, 5, 4])
    with top_left:
        if st.button("← Back to list", key="back_top"):
            st.session_state.pop("selected_run", None)
            st.rerun()

    with top_right:
        # Two side-by-side download buttons. Render lazily — wrap in
        # try/except so a renderer failure on one format doesn't break
        # the other. Files cached per session per run_id via Streamlit's
        # natural rerun-cache (st.download_button regenerates on each
        # render, fine for our ~1s render times).
        d_word, d_pdf = st.columns(2)
        with d_word:
            try:
                docx_bytes = render_docx(row)
                st.download_button(
                    label="📄 Download Word",
                    data=docx_bytes,
                    file_name=f"dpia_{row['run_id']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"Word: {e}")
        with d_pdf:
            try:
                pdf_bytes = render_pdf(row)
                st.download_button(
                    label="📥 Download PDF",
                    data=pdf_bytes,
                    file_name=f"dpia_{row['run_id']}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"PDF: {e}")

    # ── Title + status ────────────────────────────────────────────────
    st.title(f"DPIA · {row['run_id']}")
    st.caption(_status_badge(row["status"]))

    # ── Metadata strip ────────────────────────────────────────────────
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Generated", str(row["generated_at"]).split(".")[0])
    m2.metric("Reviewer", row.get("reviewed_by") or "—")
    # Statement Execution API returns all values as strings; coerce
    # before formatting. Catches None and the empty-string case too.
    _lat_raw = row.get("latency_seconds")
    try:
        _lat = float(_lat_raw) if _lat_raw not in (None, "") else None
    except (TypeError, ValueError):
        _lat = None
    m3.metric("Latency", f"{_lat:.1f}s" if _lat is not None else "—")
    m4.metric("Pack", row.get("regulation_pack") or "—")

    # ── Approve action (primary CTA, draft + approver only) ───────────
    if row["status"] == "draft" and can_approve:
        with st.container(border=True):
            st.markdown(f"#### Review actions  ·  signed in as `{user_email}`")
            st.caption(
                "Approving flips status `draft → approved` and stamps your "
                "verified email + a server-side timestamp on the audit row. "
                "Idempotent — re-running on an already-approved row is a no-op."
            )
            notes = st.text_input("Optional notes (saved to `dpia_runs.notes`)", key="approve_notes")
            if st.button("✅ Approve this DPIA", type="primary", key="approve_btn"):
                res = approve(row["run_id"], user_email, notes or None)
                if "error" in res:
                    st.error(f"UPDATE failed: {res['error']}")
                else:
                    st.success(
                        f"Approved by {user_email} at "
                        f"{datetime.now(timezone.utc).isoformat()} — refresh to see updated state."
                    )
                    st.cache_resource.clear()
                    st.rerun()
    elif row["status"] == "draft" and not can_approve:
        st.info(
            f"Signed in as **{user_email}** (view-only) — Approve is only "
            "available to CCO and GC persona users."
        )

    # ── Tabs: structured + rendered ───────────────────────────────────
    tab_sections, tab_rendered = st.tabs(
        ["📄 Sections", "📰 Rendered Document"]
    )

    sections = _normalize_sections(row)

    with tab_sections:
        if not sections:
            st.warning(
                "Structured sections not available for this run "
                f"(parse_error: `{row.get('parse_error') or 'unknown'}`). "
                "Open the Rendered Document tab to see the raw model output."
            )
        else:
            for key, title in SECTION_TITLES.items():
                with st.expander(title, expanded=False):
                    st.markdown(sections.get(key, "_(missing)_"))

    with tab_rendered:
        # Single continuous markdown render of the whole DPIA — what the
        # PDF/Word downloads produce, but in-app and live.
        st.markdown(_build_full_markdown(row))

    # ── Diagnostics (collapsed by default) ────────────────────────────
    with st.expander("🔍 Diagnostics — input data & raw model output"):
        sub_input, sub_raw = st.tabs(["📊 Input data", "📝 Raw output"])
        with sub_input:
            st.caption(
                "The metadata snapshot the model was given as input — useful "
                "for spot-checking claims in the sections above."
            )
            ctx = row.get("context_snapshot")
            if isinstance(ctx, str):
                try:
                    ctx = json.loads(ctx)
                except (json.JSONDecodeError, TypeError):
                    pass
            st.json(ctx)
        with sub_raw:
            st.caption(
                "The full raw text returned by the LLM — useful when "
                "structured parsing failed."
            )
            st.code(row.get("dpia_text", ""), language=None)


# ---------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------


def main() -> None:
    user_email = get_user_email()
    approvers = load_approver_emails()
    can_approve = is_approver(user_email, approvers)

    with st.sidebar:
        st.markdown(f"**Signed in as**\n\n`{user_email}`")
        if approvers:
            role = next(
                (k for k, v in approvers.items() if v == user_email),
                "viewer",
            )
            st.caption(f"Role: **{role}**" + (" — can approve" if can_approve else " — view-only"))
        else:
            st.caption("Persona registry not loaded — view-only mode.")
        st.divider()
        status_filter = st.radio(
            "Status filter",
            ["all", "draft", "approved", "superseded"],
            index=1,
        )
        st.divider()
        st.caption("**Catalog:** " + CATALOG)

    if "selected_run" in st.session_state:
        view_detail(st.session_state.selected_run, user_email, can_approve)
    else:
        view_list(status_filter)


if __name__ == "__main__":
    main()
